"""
RIMA ara rapor DOCX olusturucu.

Canonical source: ARA_RAPOR_RIMA.md
Output: ARA_RAPOR_RIMA.docx

Bu script eski sabit metinli v2 rapor ureticisinin yerini alir. Rapor
icerigi Markdown dosyasindan okunur; karakter anchor gorselleri
Characters/anchors altindan otomatik eklenir.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SOURCE_MD = BASE / "ARA_RAPOR_RIMA.md"
ANCHORS = BASE / "Characters" / "anchors"
OUT = BASE / "ARA_RAPOR_RIMA.docx"

CLASS_ANCHORS = [
    ("Warblade", "warblade_anchor.png"),
    ("Shadowblade", "shadowblade_anchor.png"),
    ("Ranger", "ranger_anchor.png"),
    ("Ronin", "ronin_anchor.png"),
    ("Gunslinger", "gunslinger_anchor.png"),
    ("Brawler", "brawler_anchor.png"),
    ("Ravager", "ravager_anchor.png"),
    ("Elementalist", "elementalist_anchor.png"),
    ("Hexer", "hexer_anchor.png"),
    ("Summoner", "summoner_anchor.png"),
]

C_DARK = RGBColor(0x1A, 0x1A, 0x2E)
C_ACCENT = RGBColor(0x4A, 0x6F, 0xC7)
C_BODY = RGBColor(0x21, 0x21, 0x21)
C_MUTED = RGBColor(0x66, 0x66, 0x66)
C_TABLE = "E9EDF7"


def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.size = Pt(9)
    run.font.color.rgb = C_DARK if bold else C_BODY
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_document_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = C_DARK
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def _centered(doc: Document, text: str, size: int, bold: bool = False, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_DARK


def add_title(doc: Document, meta: list[str]) -> None:
    logo_path = BASE / "STAGING" / "ornek_image1.jpeg"

    # Logo
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after = Pt(6)
    if logo_path.exists():
        logo_para.add_run().add_picture(str(logo_path), width=Inches(2.3))

    # University info
    _centered(doc, "KTO KARATAY UNIVERSITY", 16, bold=True)
    _centered(doc, "FACULTY OF ENGINEERING", 16, bold=True)
    _centered(doc, "COMPUTER ENGINEERING", 16, bold=True)

    doc.add_paragraph()

    # Report type
    _centered(doc, "PROJE RAPORU", 20, bold=True)
    _centered(doc, "Ara Rapor", 13, color=C_MUTED)

    # Spacers
    for _ in range(5):
        doc.add_paragraph()

    # Student info — tab aligned
    info_rows = [
        ("İsim ve Soyisim", "Yasin Derya Bilgin"),
        ("Öğrenci Numarası", "231450075"),
        ("Bölüm", "Bilgisayar Mühendisliği"),
        ("Proje İsmi", "RIMA - 2D Roguelite Aksiyon Oyunu"),
        ("Tarih", "Mayıs 2026"),
    ]
    for label, value in info_rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        lr = p.add_run(f"{label}\t\t: ")
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = C_DARK
        vr = p.add_run(value)
        vr.font.size = Pt(11)
        vr.font.color.rgb = C_BODY

    doc.add_page_break()


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.08
    write_inline_runs(paragraph, text)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    write_inline_runs(paragraph, text)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(3)
    write_inline_runs(paragraph, text)


def write_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text.strip())
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") else part[1:-1] if part.startswith("`") else part)
        run.font.size = Pt(11)
        run.font.color.rgb = C_BODY
        if part.startswith("**"):
            run.bold = True
        if part.startswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_index, row_values in enumerate(rows):
        for col_index, value in enumerate(row_values):
            cell = table.rows[row_index].cells[col_index]
            set_cell_text(cell, value, bold=row_index == 0)
            if row_index == 0:
                set_cell_bg(cell, C_TABLE)

    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip()
        values = [cell.strip() for cell in raw.strip("|").split("|")]
        if not all(set(cell) <= {"-", ":"} for cell in values):
            rows.append(values)
        index += 1
    return rows, index


def add_anchor_grid(doc: Document) -> None:
    available = [(name, ANCHORS / filename) for name, filename in CLASS_ANCHORS if (ANCHORS / filename).exists()]
    if not available:
        note = doc.add_paragraph()
        run = note.add_run("Not: Characters/anchors klasoru bulunamadigi icin anchor gorselleri eklenemedi.")
        run.italic = True
        run.font.color.rgb = C_MUTED
        return

    heading = doc.add_paragraph()
    run = heading.add_run("Sınıf Anchor Görselleri")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = C_DARK

    for offset in range(0, len(available), 5):
        chunk = available[offset : offset + 5]
        table = doc.add_table(rows=2, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for col_index in range(5):
            image_cell = table.rows[0].cells[col_index]
            label_cell = table.rows[1].cells[col_index]
            image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if col_index >= len(chunk):
                continue

            name, path = chunk[col_index]
            paragraph = image_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(str(path), width=Inches(0.82))
            set_cell_text(label_cell, name, bold=True)
            label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()


def add_footer(doc: Document) -> None:
    pass


def add_toc(doc: Document) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run("İçindekiler")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_DARK
    heading.paragraph_format.space_after = Pt(12)

    paragraph = doc.add_paragraph()

    run1 = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    fldChar.set(qn("w:dirty"), "true")
    run1._r.append(fldChar)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._r.append(fldChar2)

    run4 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run4._r.append(fldChar3)

    doc.add_page_break()


def build_docx() -> None:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Markdown kaynak bulunamadi: {SOURCE_MD}")

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_style(doc)

    meta: list[str] = []
    index = 0
    while index < len(lines):
        line = lines[index].strip()
        if line == "---":
            index += 1
            break
        if line.startswith("**"):
            meta.append(line)
        index += 1

    add_title(doc, meta)
    add_toc(doc)

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_markdown_table(doc, rows)
            continue

        if stripped.startswith("### "):
            title = stripped[4:].strip()
            doc.add_heading(title, level=3)
            if title == "4.1 Karakter Görsel Kimliği":
                index += 1
                continue
            index += 1
            continue

        if stripped.startswith("## "):
            title = stripped[3:].strip()
            doc.add_page_break()
            doc.add_heading(title, level=1)
            index += 1
            continue

        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            add_numbered(doc, numbered.group(1))
            index += 1
            continue

        add_paragraph(doc, stripped)
        if stripped.startswith("Her sınıf için öncelikle bir referans sprite"):
            add_anchor_grid(doc)
        index += 1

    add_footer(doc)
    doc.save(OUT)
    print(f"Olusturuldu: {OUT}")


if __name__ == "__main__":
    build_docx()
