"""
RIMA Sunum Raporu DOCX olusturucu (2026-06-18).

Sablon: create_rapor_docx.py (ayni python-docx stili).
Source:  STAGING/report/SUNUM_RAPORU_2026-06-18.md
Output:  STAGING/report/SUNUM_RAPORU_2026-06-18.docx

Markdown'da figurler inline image syntax'i ile yazilir:
    ![alt](figures_2026-06-18/xxx.png)
hemen ardindan italik altyazi satiri (*Sekil N: ...*) gelir.
Bu betik image satirini bulur, ~15cm genislikte gomer; takip eden
italik altyazi satiri ortalanmis caption olarak basilir.

Yeniden calistirma guvenlidir.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE = Path(__file__).resolve().parent
SOURCE_MD = BASE / "SUNUM_RAPORU_2026-06-18.md"
OUT = BASE / "SUNUM_RAPORU_2026-06-18.docx"

# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_ACCENT = RGBColor(0x4A, 0x6F, 0xC7)
C_BODY   = RGBColor(0x21, 0x21, 0x21)
C_MUTED  = RGBColor(0x66, 0x66, 0x66)
C_TABLE  = "E9EDF7"
C_QUOTE  = "F0F3FA"

# ---------------------------------------------------------------------------
# Markdown patterns
# ---------------------------------------------------------------------------
IMG_RE     = re.compile(r"^!\[(.*?)\]\((.+?)\)\s*$")
CAPTION_RE = re.compile(r"^\*(Şekil[^*]+)\*\s*$")


# ---------------------------------------------------------------------------
# Helpers: style
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_document_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = C_DARK
        try:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Helpers: title page
# ---------------------------------------------------------------------------

def _centered(doc: Document, text: str, size: int,
               bold: bool = False, italic: bool = False,
               color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_DARK


def _add_title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()

    _centered(doc, "RIMA — Rift Avcıları", 30, bold=True)

    doc.add_paragraph()

    _centered(doc, "Bitirme Projesi Sunum Raporu", 18, bold=True, color=C_ACCENT)
    _centered(doc, "18 Haziran 2026", 13, color=C_BODY)

    for _ in range(3):
        doc.add_paragraph()

    _centered(
        doc,
        "Bir oyun değil; bir environment + vertical slice + "
        "oyun-içi geliştirici araç seti.",
        13, italic=True, color=C_MUTED,
    )

    for _ in range(6):
        doc.add_paragraph()

    _centered(doc, "Yasin Derya Bilgin", 12, bold=True)
    _centered(doc, "RIMA — 2D Roguelite Aksiyon Oyunu", 11, color=C_BODY)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Helpers: TOC field
# ---------------------------------------------------------------------------

def _add_toc(doc: Document) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run("İçindekiler")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_DARK
    heading.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()

    r1 = p.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    fc1.set(qn("w:dirty"), "true")
    r1._r.append(fc1)

    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(instr)

    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r3._r.append(fc2)

    r4 = p.add_run()
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r4._r.append(fc3)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Helpers: inline markup
# ---------------------------------------------------------------------------

def _write_inline(paragraph, text: str) -> None:
    """Parse **bold**, *italic*, `code` in one pass."""
    token_re = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+`)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.size = Pt(11)
        run.font.color.rgb = C_BODY
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run.text = part


# ---------------------------------------------------------------------------
# Helpers: body paragraph
# ---------------------------------------------------------------------------

def _add_paragraph(doc: Document, text: str, bullet: bool = False) -> None:
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    _write_inline(p, text)


def _add_quote(doc: Document, text: str) -> None:
    """Blockquote (>) -> shaded, accent-bordered paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), C_QUOTE)
    pPr.append(shd)
    _write_inline(p, text)
    for run in p.runs:
        run.italic = True


# ---------------------------------------------------------------------------
# Helpers: table
# ---------------------------------------------------------------------------

def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r"[-:]+", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_values in enumerate(rows):
        for ci in range(n_cols):
            cell = table.rows[ri].cells[ci]
            value = row_values[ci] if ci < len(row_values) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            _write_inline(p, value)
            for run in p.runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = C_DARK
                else:
                    run.font.color.rgb = C_BODY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                _set_cell_bg(cell, C_TABLE)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Helpers: figure (inline markdown image + following italic caption)
# ---------------------------------------------------------------------------

def _add_figure(doc: Document, img_rel: str, caption: str | None) -> bool:
    img_path = (BASE / img_rel).resolve()
    embedded = False
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run()
        run.add_picture(str(img_path), width=Cm(15))
        embedded = True
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"[Görsel bulunamadı: {img_rel}]")
        run.italic = True
        run.font.color.rgb = C_MUTED

    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(10)
        cap_run = cap_p.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        cap_run.font.color.rgb = C_MUTED

    return embedded


# ---------------------------------------------------------------------------
# Helpers: code block
# ---------------------------------------------------------------------------

def _add_code_block(doc: Document, lines: list[str]) -> None:
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = C_BODY


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_docx() -> int:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Kaynak bulunamadi: {SOURCE_MD}")

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()

    doc = Document()
    _set_document_style(doc)

    _add_title_page(doc)
    _add_toc(doc)

    img_embedded = 0
    i = 0
    n = len(lines)

    # Cover block (title + subtitle + tagline before "## 1.") is rendered
    # by _add_title_page; skip it in the body to avoid duplication.
    body_start_re = re.compile(r"^##\s+\d+\.")
    while i < n and not body_start_re.match(lines[i].strip()):
        i += 1

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line / HR
        if not stripped or stripped == "---":
            i += 1
            continue

        # Document-level "# Title" -> skip (cover already done)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        # Heading 1: ##
        if stripped.startswith("## ") and not stripped.startswith("### "):
            title = stripped[3:].strip()
            doc.add_page_break()
            doc.add_heading(title, level=1)
            i += 1
            continue

        # Heading 2: ###
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            title = stripped[4:].strip()
            doc.add_heading(title, level=2)
            i += 1
            continue

        # Heading 3: ####
        if stripped.startswith("#### "):
            title = stripped[5:].strip()
            doc.add_heading(title, level=3)
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            _add_code_block(doc, code_lines)
            continue

        # Table
        if stripped.startswith("|"):
            rows, i = _parse_table(lines, i)
            _add_table(doc, rows)
            continue

        # Inline markdown image: ![alt](path) + optional following caption
        img_match = IMG_RE.match(stripped)
        if img_match:
            img_rel = img_match.group(2).strip()
            caption = None
            j = i + 1
            # skip a single blank line between image and caption
            while j < n and not lines[j].strip():
                j += 1
            if j < n:
                cap_match = CAPTION_RE.match(lines[j].strip())
                if cap_match:
                    caption = cap_match.group(1).strip()
                    i = j  # consume the caption line too
            if _add_figure(doc, img_rel, caption):
                img_embedded += 1
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            _add_quote(doc, stripped[2:].strip())
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            _add_paragraph(doc, stripped[2:], bullet=True)
            i += 1
            continue

        # Numbered list
        num_m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_m:
            _add_paragraph(doc, num_m.group(1), bullet=True)
            i += 1
            continue

        # Plain paragraph
        _add_paragraph(doc, stripped)
        i += 1

    doc.save(OUT)
    print(f"Olusturuldu: {OUT}")
    return img_embedded


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------

def verify_docx(expected_imgs: int) -> None:
    if not OUT.exists():
        print("HATA: docx bulunamadi.")
        return

    size_kb = OUT.stat().st_size / 1024.0
    doc = Document(str(OUT))
    para_count = len(doc.paragraphs)

    h1_list = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]

    img_count = len(doc.element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    ))

    print("\n=== Dogrulama ===")
    print(f"Dosya boyutu    : {size_kb:.1f} KB  ({'OK >20KB' if size_kb > 20 else 'UYARI <20KB'})")
    print(f"Toplam paragraf : {para_count}")
    print(f"Heading 1 sayisi: {len(h1_list)}")
    for h in h1_list:
        print(f"  - {h}")
    print(f"Gomulu gorsel   : {img_count}  (build sirasinda gomulen: {expected_imgs})")


# ---------------------------------------------------------------------------

if __name__ == "__main__":
    embedded = build_docx()
    verify_docx(embedded)
