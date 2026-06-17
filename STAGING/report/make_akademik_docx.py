"""
RIMA Senior Design Report DOCX olusturucu (KTO Karatay formati).

Template: create_rapor_docx.py (BarberApp stili).
Source:   STAGING/report/RIMA_Senior_Design_Report.md
Output:   STAGING/report/RIMA_Senior_Design_Report.docx

Kapak referansi:  F:/BarberApp/Yasin_Derya_Bilgin_Senior_Design_1.pdf
  - Ust-orta KTO logo
  - KTO KARATAY UNIVERSITY / FACULTY OF ENGINEERING / COMPUTER ENGINEERING
  - SENIOR DESIGN PROJECT - 2
  - Sol-alt bilgi blogu (Isim / Ogrenci No / Proje)
  - Numarali teal basliklar, Icindekiler (TOC field), monospace kod, sayfa no.

Yeniden calistirma guvenlidir.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE = Path(__file__).resolve().parent
SOURCE_MD = BASE / "RIMA_Senior_Design_Report.md"
FIG_2026 = BASE / "figures_2026-06-18"
FIG_V2 = BASE / "figures_v2"
GRAPHIFY = BASE / "graphify"
SCREENSHOTS = BASE.parent / "report_screenshots"
LOGO = FIG_2026 / "kto_logo.png"
OUT = BASE / "RIMA_Senior_Design_Report.docx"

# ---------------------------------------------------------------------------
# Colors  (referans teal/koyu-mavi tonu)
# ---------------------------------------------------------------------------
C_TEAL   = RGBColor(0x21, 0x57, 0x6B)  # numarali basliklar (~#21576B)
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_ACCENT = RGBColor(0x21, 0x57, 0x6B)
C_BODY   = RGBColor(0x21, 0x21, 0x21)
C_MUTED  = RGBColor(0x66, 0x66, 0x66)
C_TABLE  = "E3ECF0"  # acik teal-gri baslik hucresi

# ---------------------------------------------------------------------------
# Figure -> file resolution.
# Brief figure caption suffix "| <path>" verir; path birden cok klasore gore
# cozulur. Asagidaki harita yedek cozumdur (caption'da path yoksa).
# ---------------------------------------------------------------------------
FIGURE_DIRS = [BASE, FIG_2026, FIG_V2, GRAPHIFY, SCREENSHOTS, BASE.parent]


def _resolve_figure(path_hint: str) -> Path | None:
    """Caption sonrasi '| dosya' ipucunu gercek dosya yoluna cevirir."""
    hint = path_hint.strip()
    if not hint:
        return None
    # Mutlak / proje-koku yollar
    candidates: list[Path] = []
    # 1) figures_2026-06-18 oncelikli: cıplak ad
    bare = Path(hint).name
    # 2) 'figures_v2/...' veya 'report_screenshots/...' gibi alt-yol ipuclari
    for d in FIGURE_DIRS:
        candidates.append(d / hint)        # tam ipucu (alt-yol dahil)
        candidates.append(d / bare)        # ciplak dosya adi
    for c in candidates:
        if c.exists() and c.is_file():
            return c
    return None


# ---------------------------------------------------------------------------
# Helpers: style
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_document_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.7)
        section.right_margin  = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass

    # Teal numarali basliklar (referans tonu)
    sizes = {"Heading 1": 15, "Heading 2": 13, "Heading 3": 11.5}
    for style_name, sz in sizes.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.size = Pt(sz)
        style.font.color.rgb = C_TEAL
        try:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        except Exception:
            pass


def _add_page_number_footer(doc: Document) -> None:
    """Alt-sag PAGE field."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r1 = p.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    r1._r.append(fc1)
    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    r2._r.append(instr)
    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r3._r.append(fc2)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = C_MUTED


# ---------------------------------------------------------------------------
# Helpers: title page (referans birebir)
# ---------------------------------------------------------------------------

def _centered(doc: Document, text: str, size: float,
              bold: bool = False, color: RGBColor | None = None,
              space_after: int = 2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_DARK


def _add_title_page(doc: Document) -> None:
    # bir miktar ust bosluk
    for _ in range(3):
        doc.add_paragraph()

    # Logo (ortali)
    if LOGO.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_after = Pt(14)
        lp.add_run().add_picture(str(LOGO), width=Cm(6.0))

    _centered(doc, "KTO KARATAY UNIVERSITY", 14, bold=True)
    _centered(doc, "FACULTY OF ENGINEERING", 14, bold=True)
    _centered(doc, "COMPUTER ENGİNEERİNG",   14, bold=True)

    for _ in range(3):
        doc.add_paragraph()

    _centered(doc, "SENIOR DESIGN PROJECT - 2", 14, bold=True)

    for _ in range(8):
        doc.add_paragraph()

    info_rows = [
        ("İsim ve Soyisim",  "Yasin Derya Bilgin"),
        ("Öğrenci Numarası", "231450075"),
        ("Proje İsmi",       "RIMA - Rift Avcıları"),
    ]
    for label, value in info_rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        lr = p.add_run(f"{label}".ljust(20))
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = C_DARK
        sep = p.add_run(":  ")
        sep.font.size = Pt(11)
        sep.font.color.rgb = C_DARK
        vr = p.add_run(value)
        vr.font.size = Pt(11)
        vr.font.color.rgb = C_BODY

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Helpers: TOC field  (Word'de F9 ile guncellenir)
# ---------------------------------------------------------------------------

def _add_toc(doc: Document) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run("İçindekiler")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = C_TEAL
    heading.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    r1 = p.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    fc1.set(qn("w:dirty"), "true")
    r1._r.append(fc1)

    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(instr)

    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r3._r.append(fc2)

    # Word TOC'u guncelleyene kadar gosterilecek yer-tutucu
    r_ph = p.add_run("(Word'de güncellemek için: tabloya sağ tık → Alanı Güncelle / F9)")
    r_ph.italic = True
    r_ph.font.size = Pt(9)
    r_ph.font.color.rgb = C_MUTED

    r4 = p.add_run()
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r4._r.append(fc3)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Helpers: inline markup
# ---------------------------------------------------------------------------

def _write_inline(paragraph, text: str) -> None:
    token_re = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+`)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.size = Pt(11)
        run.font.color.rgb = C_BODY
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run.text = part


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    _write_inline(p, text)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    _write_inline(p, text)


# ---------------------------------------------------------------------------
# Helpers: table
# ---------------------------------------------------------------------------

def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r"[-:]+", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_values in enumerate(rows):
        for ci in range(n_cols):
            cell = table.rows[ri].cells[ci]
            value = row_values[ci] if ci < len(row_values) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            _write_inline(p, value)
            for run in p.runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = C_TEAL
                else:
                    run.font.color.rgb = C_BODY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                _set_cell_bg(cell, C_TABLE)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Helpers: figure
# ---------------------------------------------------------------------------

def _add_figure(doc: Document, fig_num: int, caption: str, path_hint: str,
                embedded_counter: list[int]) -> None:
    img_path = _resolve_figure(path_hint)
    if img_path is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        p.add_run().add_picture(str(img_path), width=Cm(14.5))
        embedded_counter[0] += 1
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"[Şekil {fig_num} — görsel bulunamadı: {path_hint}]")
        run.italic = True
        run.font.color.rgb = C_MUTED

    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(10)
    cap_run = cap_p.add_run(f"Şekil {fig_num}: {caption}")
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = C_MUTED


# ---------------------------------------------------------------------------
# Helpers: code block
# ---------------------------------------------------------------------------

def _add_code_block(doc: Document, code_lines: list[str]) -> None:
    text = "\n".join(code_lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = C_BODY


# ---------------------------------------------------------------------------
# Figure line pattern:  [Şekil N: caption | path_hint]
# ---------------------------------------------------------------------------
FIGURE_RE = re.compile(r"^\[Şekil\s+(\d+)\s*:\s*(.+?)\s*\|\s*(.+?)\]\s*$")


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_docx() -> int:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Kaynak bulunamadi: {SOURCE_MD}")

    raw_lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()

    # HTML yorumlarini at
    filtered: list[str] = []
    in_html_comment = False
    for line in raw_lines:
        if "<!--" in line and "-->" in line:
            clean = re.sub(r"<!--.*?-->", "", line).rstrip()
            if clean.strip():
                filtered.append(clean)
            continue
        if "<!--" in line:
            in_html_comment = True
            before = line[:line.index("<!--")].rstrip()
            if before.strip():
                filtered.append(before)
            continue
        if "-->" in line:
            in_html_comment = False
            after = line[line.index("-->") + 3:].rstrip()
            if after.strip():
                filtered.append(after)
            continue
        if in_html_comment:
            continue
        filtered.append(line)

    lines = filtered
    doc = Document()
    _set_document_style(doc)

    _add_title_page(doc)
    _add_toc(doc)
    _add_page_number_footer(doc)

    embedded_counter = [0]

    i = 0
    n = len(lines)
    first_h1_done = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        # Belge-duzeyi `# Baslik` -> atla
        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        # Heading 1: ##
        if stripped.startswith("## ") and not stripped.startswith("### "):
            title = stripped[3:].strip()
            if first_h1_done:
                doc.add_page_break()
            first_h1_done = True
            doc.add_heading(title, level=1)
            i += 1
            continue

        # Heading 2: ###
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue

        # Heading 3: ####
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1
            _add_code_block(doc, code_lines)
            continue

        # Table
        if stripped.startswith("|"):
            rows, i = _parse_table(lines, i)
            _add_table(doc, rows)
            continue

        # Figure
        fig_match = FIGURE_RE.match(stripped)
        if fig_match:
            fig_num = int(fig_match.group(1))
            caption = fig_match.group(2).strip()
            hint = fig_match.group(3).strip()
            _add_figure(doc, fig_num, caption, hint, embedded_counter)
            i += 1
            continue

        # Bullet
        if stripped.startswith("- "):
            _add_bullet(doc, stripped[2:])
            i += 1
            continue

        # Numbered list
        num_m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_m:
            _add_bullet(doc, num_m.group(1))
            i += 1
            continue

        _add_paragraph(doc, stripped)
        i += 1

    doc.save(OUT)
    print(f"Olusturuldu: {OUT}")
    return embedded_counter[0]


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------

def verify_docx(expected_imgs: int) -> None:
    if not OUT.exists():
        print("HATA: docx bulunamadi.")
        return

    size_kb = OUT.stat().st_size / 1024
    doc = Document(str(OUT))
    para_count = len(doc.paragraphs)

    h1_list = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    img_count = len(doc.element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    ))

    print("\n=== Dogrulama ===")
    print(f"Dosya boyutu    : {size_kb:.0f} KB  ({'OK >100KB' if size_kb > 100 else 'UYARI <100KB'})")
    print(f"Toplam paragraf : {para_count}")
    print(f"Heading 1 sayisi: {len(h1_list)}")
    for h in h1_list:
        print(f"  - {h}")
    print(f"Gomulu gorsel   : {img_count}  (gomulen figur sayaci: {expected_imgs} + kapak logo)")


# ---------------------------------------------------------------------------

if __name__ == "__main__":
    embedded = build_docx()
    verify_docx(embedded)
