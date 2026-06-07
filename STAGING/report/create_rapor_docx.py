"""
RIMA Final Rapor DOCX olusturucu.

Source:  STAGING/report/RAPOR_DRAFT_2026-06-06.md
Output:  STAGING/report/RAPOR_RIMA_2026-06-06.docx

Yeniden calistirma guvenlidir: calistirildikca docx'i yazar.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE = Path(__file__).resolve().parent
SOURCE_MD = BASE / "RAPOR_DRAFT_2026-06-06.md"
SCREENSHOTS = BASE.parent / "report_screenshots"
OUT = BASE / "RAPOR_RIMA_2026-06-06.docx"

# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_ACCENT = RGBColor(0x4A, 0x6F, 0xC7)
C_BODY   = RGBColor(0x21, 0x21, 0x21)
C_MUTED  = RGBColor(0x66, 0x66, 0x66)
C_TABLE  = "E9EDF7"

# ---------------------------------------------------------------------------
# Figure → file mapping  (figure number → filename relative to SCREENSHOTS)
# ---------------------------------------------------------------------------
FIGURE_FILES: dict[int, str | None] = {
    1:  "02_chamber_overview.png",
    2:  "03_chamber_g_prompt.png",
    3:  "05_combat.png",
    4:  "06_draft_cards.png",
    5:  "07_branching_doors.png",
    6:  None,   # no file
    7:  "04_run_room_overview.png",
    8:  None,   # no file
    9:  "11_map_designer.png",
    10: "12_room_browser_scene.png",
    11: None,   # no file
    12: None,   # placeholder — pedestal close-up to be recaptured
    13: None,   # no file
    14: None,   # no file
}

# ---------------------------------------------------------------------------
# Helpers: style
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_document_style(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin    = Cm(2.4)
        section.bottom_margin = Cm(2.4)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    try:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    except Exception:
        pass

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = C_DARK
        try:
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Helpers: title page
# ---------------------------------------------------------------------------

def _centered(doc: Document, text: str, size: int,
               bold: bool = False, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or C_DARK


def _add_title_page(doc: Document) -> None:
    # University block
    _centered(doc, "KTO KARATAY ÜNİVERSİTESİ", 16, bold=True)
    _centered(doc, "MÜHENDİSLİK FAKÜLTESİ",   16, bold=True)
    _centered(doc, "BİLGİSAYAR MÜHENDİSLİĞİ", 16, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    _centered(doc, "BİTİRME PROJESİ RAPORU", 22, bold=True)
    _centered(doc, "RIMA — Bitirme Projesi Final Raporu", 14, color=C_ACCENT)

    for _ in range(6):
        doc.add_paragraph()

    info_rows = [
        ("İsim ve Soyisim",   "Yasin Derya Bilgin"),
        ("Öğrenci Numarası",  "231450075"),
        ("Bölüm",             "Bilgisayar Mühendisliği"),
        ("Proje İsmi",        "RIMA — 2D Roguelite Aksiyon Oyunu"),
        ("Tarih",             "2026-06-06"),
    ]
    for label, value in info_rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        lr = p.add_run(f"{label}\t\t: ")
        lr.bold = True
        lr.font.size = Pt(11)
        lr.font.color.rgb = C_DARK
        vr = p.add_run(value)
        vr.font.size = Pt(11)
        vr.font.color.rgb = C_BODY

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Helpers: TOC field
# ---------------------------------------------------------------------------

def _add_toc(doc: Document) -> None:
    heading = doc.add_paragraph()
    run = heading.add_run("İçindekiler")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = C_DARK
    heading.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()

    r1 = p.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    fc1.set(qn("w:dirty"), "true")
    r1._r.append(fc1)

    r2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(instr)

    r3 = p.add_run()
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "separate")
    r3._r.append(fc2)

    r4 = p.add_run()
    fc3 = OxmlElement("w:fldChar")
    fc3.set(qn("w:fldCharType"), "end")
    r4._r.append(fc3)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Helpers: inline markup
# ---------------------------------------------------------------------------

def _write_inline(paragraph, text: str) -> None:
    """Parse **bold**, *italic*, `code` in one pass."""
    # Combined pattern: **bold**, *italic*, `code`
    token_re = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+`)")
    parts = token_re.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.size = Pt(11)
        run.font.color.rgb = C_BODY
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            run.text = part


# ---------------------------------------------------------------------------
# Helpers: body paragraph
# ---------------------------------------------------------------------------

def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    _write_inline(p, text)


# ---------------------------------------------------------------------------
# Helpers: table
# ---------------------------------------------------------------------------

def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        # Skip separator rows (only dashes/colons)
        if not all(re.fullmatch(r"[-:]+", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, row_values in enumerate(rows):
        for ci in range(n_cols):
            cell = table.rows[ri].cells[ci]
            value = row_values[ci] if ci < len(row_values) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            _write_inline(p, value)
            for run in p.runs:
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = C_DARK
                else:
                    run.font.color.rgb = C_BODY
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                _set_cell_bg(cell, C_TABLE)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Helpers: figures
# ---------------------------------------------------------------------------

def _add_figure(doc: Document, fig_num: int, caption: str) -> None:
    filename = FIGURE_FILES.get(fig_num)
    if filename is not None:
        img_path = SCREENSHOTS / filename
        if img_path.exists():
            # Image paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run()
            run.add_picture(str(img_path), width=Cm(15))
        else:
            # File listed but missing — treat as placeholder
            filename = None

    if filename is None or not (SCREENSHOTS / (filename or "")).exists():
        # Placeholder text
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"[Şekil {fig_num} — görsel eklenecek: {caption}]")
        run.italic = True
        run.font.color.rgb = C_MUTED

    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_after = Pt(10)
    cap_run = cap_p.add_run(f"Şekil {fig_num}: {caption}")
    cap_run.italic = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = C_MUTED


# ---------------------------------------------------------------------------
# Helpers: code block
# ---------------------------------------------------------------------------

def _add_code_block(doc: Document, lines: list[str]) -> None:
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = C_BODY


# ---------------------------------------------------------------------------
# Figure line pattern
# ---------------------------------------------------------------------------
FIGURE_RE = re.compile(r"^\[Şekil\s+(\d+)\s*:\s*(.+?)\]\s*$")


# ---------------------------------------------------------------------------
# Main builder
# ---------------------------------------------------------------------------

def build_docx() -> None:
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Kaynak bulunamadi: {SOURCE_MD}")

    raw_lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()

    # ---- Pre-process: strip HTML comments and word-count lines ----------
    filtered: list[str] = []
    in_html_comment = False
    for line in raw_lines:
        if "<!--" in line and "-->" in line:
            # inline comment — remove it but keep rest of line
            clean = re.sub(r"<!--.*?-->", "", line).rstrip()
            if clean.strip():
                filtered.append(clean)
            continue
        if "<!--" in line:
            in_html_comment = True
            before = line[:line.index("<!--")].rstrip()
            if before.strip():
                filtered.append(before)
            continue
        if "-->" in line:
            in_html_comment = False
            after = line[line.index("-->") + 3:].rstrip()
            if after.strip():
                filtered.append(after)
            continue
        if in_html_comment:
            continue
        # Strip *Kelime sayısı: ...* lines
        if re.match(r"^\*Kelime sayısı", line.strip()):
            continue
        filtered.append(line)

    lines = filtered
    doc = Document()
    _set_document_style(doc)

    # ---- Title + TOC ----------------------------------------------------
    _add_title_page(doc)
    _add_toc(doc)

    # ---- Parse lines -----------------------------------------------------
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line / HR
        if not stripped or stripped == "---":
            i += 1
            continue

        # Skip document-level `# Title` (the very first heading)
        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        # Heading 1: ##
        if stripped.startswith("## ") and not stripped.startswith("### "):
            title = stripped[3:].strip()
            doc.add_page_break()
            h = doc.add_heading(title, level=1)
            i += 1
            continue

        # Heading 2: ###
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            title = stripped[4:].strip()
            h = doc.add_heading(title, level=2)
            i += 1
            continue

        # Heading 3: ####
        if stripped.startswith("#### "):
            title = stripped[5:].strip()
            h = doc.add_heading(title, level=3)
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1  # closing ```
            _add_code_block(doc, code_lines)
            continue

        # Table
        if stripped.startswith("|"):
            rows, i = _parse_table(lines, i)
            _add_table(doc, rows)
            continue

        # Figure line: [Şekil N: caption]
        fig_match = FIGURE_RE.match(stripped)
        if fig_match:
            fig_num = int(fig_match.group(1))
            caption = fig_match.group(2).strip()
            _add_figure(doc, fig_num, caption)
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            _add_paragraph(doc, stripped[2:])
            i += 1
            continue

        # Numbered list
        num_m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_m:
            _add_paragraph(doc, num_m.group(1))
            i += 1
            continue

        # Italic-only "savunma notu" lines like *[...]*
        if stripped.startswith("*[") and stripped.endswith("]*"):
            _add_paragraph(doc, stripped)
            i += 1
            continue

        # Plain paragraph
        _add_paragraph(doc, stripped)
        i += 1

    doc.save(OUT)
    print(f"Olusturuldu: {OUT}")


# ---------------------------------------------------------------------------
# Verification
# ---------------------------------------------------------------------------

def verify_docx() -> None:
    if not OUT.exists():
        print("HATA: docx bulunamadi.")
        return

    doc = Document(str(OUT))
    para_count = len(doc.paragraphs)

    h1_list: list[str] = []
    for p in doc.paragraphs:
        if p.style.name == "Heading 1":
            h1_list.append(p.text)

    # Count inline images (blip elements)
    img_count = len(doc.element.findall(
        ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    ))

    print(f"\n=== Dogrulama ===")
    print(f"Toplam paragraf : {para_count}")
    print(f"Heading 1 sayisi: {len(h1_list)}")
    for h in h1_list:
        print(f"  - {h}")
    print(f"Gomulu gorsel   : {img_count}  (beklenen: 8)")
    if img_count != 8:
        print("  UYARI: gorsel sayisi 8 degil!")
    else:
        print("  OK")


# ---------------------------------------------------------------------------

if __name__ == "__main__":
    build_docx()
    verify_docx()
